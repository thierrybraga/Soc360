import os
import json
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_full_report(org_id, org_name, period_start, period_end, report_data, output_dir, pdf_output_dir=None):
    os.makedirs(output_dir, exist_ok=True)
    if pdf_output_dir:
        os.makedirs(pdf_output_dir, exist_ok=True)

    docx_filename = f"relatorio_umbrela_{org_id}_{period_start}_{period_end}.docx"
    docx_path = os.path.join(output_dir, docx_filename)

    doc = Document()

    title = doc.add_heading(f'Relatório Cisco Umbrella - {org_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Período: {period_start} a {period_end}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _add_deployment_section(doc, report_data.get('deployment', {}))
    _add_activity_section(doc, report_data.get('activity', {}))
    _add_security_section(doc, report_data.get('security_categories', {}))
    _add_app_discovery_section(doc, report_data.get('app_discovery', {}))
    _add_security_requests_section(doc, report_data.get('security_requests', {}))

    doc.save(docx_path)

    result = {'docx_filename': docx_filename, 'docx': docx_path}

    if pdf_output_dir:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        pdf_path = os.path.join(pdf_output_dir, pdf_filename)
        pdf_result = _convert_to_pdf(docx_path, pdf_path)
        if pdf_result:
            result['pdf_filename'] = pdf_filename
            result['pdf'] = pdf_path

    return result


def _add_deployment_section(doc, deployment):
    doc.add_heading('1. Deployment', level=1)

    networks = deployment.get('networks', [])
    if networks:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nome'
        hdr_cells[1].text = 'IP/CIDR'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Dinâmica'
        hdr_cells[4].text = 'Política'
        for net in networks:
            row_cells = table.add_row().cells
            row_cells[0].text = net.get('name', '')
            row_cells[1].text = net.get('ipAddress', '')
            row_cells[2].text = net.get('status', '')
            row_cells[3].text = 'Sim' if net.get('isDynamic') else 'Não'
            row_cells[4].text = net.get('primaryPolicy', '')
    else:
        doc.add_paragraph('Nenhuma rede encontrada.')

    roaming = deployment.get('roaming_computers', [])
    if roaming:
        doc.add_heading('Computadores Roaming', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nome'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'Última Sincronia'
        hdr_cells[3].text = 'SO'
        for comp in roaming[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = comp.get('name', '')
            row_cells[1].text = comp.get('status', '')
            row_cells[2].text = comp.get('lastSync', '')
            row_cells[3].text = comp.get('osVersion', '')

    vas = deployment.get('virtual_appliances', [])
    if vas:
        doc.add_heading('Virtual Appliances', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nome'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'IP'
        for va in vas:
            row_cells = table.add_row().cells
            row_cells[0].text = va.get('name', '')
            row_cells[1].text = va.get('status', '')
            row_cells[2].text = va.get('ipAddress', '')

    doc.add_paragraph()


def _add_activity_section(doc, activity):
    doc.add_heading('2. Atividade DNS', level=1)
    summary = activity.get('summary', {})
    if summary:
        doc.add_paragraph(f"Total de Requisições: {summary.get('totalRequests', 0):,}")
        doc.add_paragraph(f"Total de Bloqueios: {summary.get('totalBlocks', 0):,}")
        doc.add_paragraph(f"Bloqueios de Segurança: {summary.get('securityBlocks', 0):,}")
    else:
        doc.add_paragraph('Nenhum dado de atividade disponível.')
    doc.add_paragraph()


def _add_security_section(doc, security):
    doc.add_heading('3. Categorias de Segurança', level=1)
    categories = security.get('categories', {})
    if categories:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Categoria'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Variação %'
        for cat_name, cat_data in categories.items():
            row_cells = table.add_row().cells
            row_cells[0].text = cat_name.replace('_', ' ').title()
            row_cells[1].text = str(cat_data.get('total', 0))
            row_cells[2].text = f"{cat_data.get('percentChange', 0):+g}%"
    else:
        doc.add_paragraph('Nenhum dado de segurança disponível.')
    doc.add_paragraph()


def _add_app_discovery_section(doc, app_discovery):
    doc.add_heading('4. Descoberta de Aplicativos', level=1)
    if app_discovery:
        doc.add_paragraph(f"Total de Aplicativos: {app_discovery.get('totalApps', 0)}")
        doc.add_paragraph(f"Aplicativos de Risco: {app_discovery.get('riskyApps', 0)}")
        apps = app_discovery.get('applications', [])
        if apps:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Aplicativo'
            hdr_cells[1].text = 'Categoria'
            hdr_cells[2].text = 'Nível de Risco'
            for app in apps:
                row_cells = table.add_row().cells
                row_cells[0].text = app.get('name', '')
                row_cells[1].text = app.get('category', '')
                row_cells[2].text = app.get('riskLevel', '')
    else:
        doc.add_paragraph('Nenhum dado de aplicativo disponível.')
    doc.add_paragraph()


def _add_security_requests_section(doc, security_requests):
    doc.add_heading('5. Requisições de Segurança por Identidade', level=1)
    identities = security_requests.get('identities', [])
    if identities:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Identidade'
        hdr_cells[1].text = 'Tipo'
        hdr_cells[2].text = 'Bloqueios'
        for ident in identities[:15]:
            row_cells = table.add_row().cells
            row_cells[0].text = ident.get('name', '')
            row_cells[1].text = ident.get('type', '')
            row_cells[2].text = f"{ident.get('blockedRequests', 0):,}"
    else:
        doc.add_paragraph('Nenhum dado de identidade disponível.')
    doc.add_paragraph()


def _convert_to_pdf(docx_path, pdf_path):
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path],
            capture_output=True, timeout=60
        )
        return result.returncode == 0
    except Exception:
        return False
